from markitdown import MarkItDown
import io
import tempfile
import os

def convert_html_to_format(html_bytes: bytes, output_format: str):
    html_content = html_bytes.decode('utf-8')
    
    if output_format == "markdown":
        # Use MarkItDown for clean markdown conversion
        md = MarkItDown(enable_plugins=False)
        input_stream = io.BytesIO(html_bytes)
        result = md.convert_stream(input_stream, source_format="html")
        return {
            "content": result.text_content.encode("utf-8"),
            "media_type": "text/markdown",
            "filename": "conversation.md"
        }

    elif output_format == "pdf":
        # Render-optimized PDF conversion (avoid external dependencies)
        
        # Method 1: Enhanced FPDF with html2text (most reliable for Render)
        try:
            from fpdf import FPDF
            from html2text import html2text
            
            # Configure html2text for better conversion
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = True
            h.ignore_emphasis = False
            h.body_width = 80
            h.wrap_links = False
            text_content = h.handle(html_content)
            
            # Create PDF with enhanced formatting
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            
            # Process content line by line with formatting
            lines = text_content.splitlines()
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    pdf.ln(3)
                    continue
                
                try:
                    # Handle different content types
                    if line.startswith('#'):
                        # Headers
                        level = len(line) - len(line.lstrip('#'))
                        header_text = line.lstrip('# ').strip()
                        if header_text:
                            font_size = max(18 - level * 2, 11)
                            pdf.set_font("Arial", 'B', font_size)
                            pdf.ln(4)
                            # Ensure text fits in Latin-1
                            safe_text = header_text.encode('latin1', 'replace').decode('latin1')
                            pdf.cell(0, 8, safe_text, ln=True)
                            pdf.ln(2)
                            pdf.set_font("Arial", size=10)
                    
                    # Handle bold text markers
                    elif '**' in line and line.count('**') >= 2:
                        pdf.set_font("Arial", 'B', 10)
                        clean_line = line.replace('**', '').encode('latin1', 'replace').decode('latin1')
                        pdf.multi_cell(0, 6, clean_line)
                        pdf.set_font("Arial", size=10)
                    
                    # Handle code blocks
                    elif line.startswith('```') or line.startswith('    ') or line.startswith('\t'):
                        if not line.startswith('```'):  # Don't print the ``` markers
                            pdf.set_font("Courier", size=9)
                            clean_line = line.lstrip(' \t`').encode('latin1', 'replace').decode('latin1')
                            if clean_line:
                                pdf.multi_cell(0, 5, clean_line)
                            pdf.set_font("Arial", size=10)
                    
                    # Handle bullet points
                    elif line.startswith('* ') or line.startswith('- ') or line.startswith('+ '):
                        pdf.set_font("Arial", size=10)
                        bullet_text = "• " + line[2:].encode('latin1', 'replace').decode('latin1')
                        pdf.multi_cell(0, 6, bullet_text)
                    
                    # Handle numbered lists
                    elif len(line) > 2 and line[0].isdigit() and line[1:3] in ['. ', ') ']:
                        pdf.set_font("Arial", size=10)
                        clean_line = line.encode('latin1', 'replace').decode('latin1')
                        pdf.multi_cell(0, 6, clean_line)
                    
                    # Handle quotes
                    elif line.startswith('>'):
                        pdf.set_font("Arial", 'I', 10)
                        quote_text = line.lstrip('> ').encode('latin1', 'replace').decode('latin1')
                        if quote_text:
                            pdf.multi_cell(0, 6, f'"{quote_text}"')
                        pdf.set_font("Arial", size=10)
                    
                    # Regular text
                    else:
                        pdf.set_font("Arial", size=10)
                        clean_line = line.encode('latin1', 'replace').decode('latin1')
                        if clean_line:
                            pdf.multi_cell(0, 6, clean_line)
                
                except Exception as line_error:
                    # Skip problematic lines but continue processing
                    print(f"Skipped line due to error: {line_error}")
                    continue
            
            pdf_content = pdf.output(dest='S').encode('latin-1')
            return {
                "content": pdf_content,
                "media_type": "application/pdf",
                "filename": "conversation.pdf"
            }
            
        except Exception as e:
            print(f"Enhanced FPDF failed: {e}")
            
            # Method 2: Basic FPDF fallback (always works on Render)
            try:
                from fpdf import FPDF
                from bs4 import BeautifulSoup
                
                # Parse HTML and extract clean text
                soup = BeautifulSoup(html_content, 'html.parser')
                
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Arial", size=10)
                
                # Process each HTML element
                for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'pre', 'code', 'blockquote']):
                    text = element.get_text().strip()
                    if not text:
                        continue
                    
                    try:
                        # Style based on element type
                        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                            level = int(element.name[1])
                            pdf.set_font("Arial", 'B', max(16 - level, 10))
                            pdf.ln(4)
                            safe_text = text.encode('latin1', 'replace').decode('latin1')
                            pdf.multi_cell(0, 8, safe_text)
                            pdf.ln(2)
                            pdf.set_font("Arial", size=10)
                        elif element.name in ['pre', 'code']:
                            pdf.set_font("Courier", size=9)
                            safe_text = text.encode('latin1', 'replace').decode('latin1')
                            pdf.multi_cell(0, 5, safe_text)
                            pdf.set_font("Arial", size=10)
                        elif element.name == 'blockquote':
                            pdf.set_font("Arial", 'I', 10)
                            safe_text = f'"{text}"'.encode('latin1', 'replace').decode('latin1')
                            pdf.multi_cell(0, 6, safe_text)
                            pdf.set_font("Arial", size=10)
                        else:
                            safe_text = text.encode('latin1', 'replace').decode('latin1')
                            pdf.multi_cell(0, 6, safe_text)
                    
                    except Exception:
                        # Skip problematic elements
                        continue
                
                pdf_content = pdf.output(dest='S').encode('latin-1')
                return {
                    "content": pdf_content,
                    "media_type": "application/pdf",
                    "filename": "conversation.pdf"
                }
            
            except Exception as final_error:
                print(f"All PDF methods failed: {final_error}")
                # Return a simple error PDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.cell(0, 10, "Error: Could not convert content to PDF", ln=True)
                pdf_content = pdf.output(dest='S').encode('latin-1')
                return {
                    "content": pdf_content,
                    "media_type": "application/pdf",
                    "filename": "error.pdf"
                }

    elif output_format == "docx":
        # Render-optimized DOCX conversion
        try:
            from docx import Document
            from bs4 import BeautifulSoup
            
            # Parse HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            doc = Document()
            
            # Add title
            title = soup.find('title')
            if title:
                doc.add_heading(title.get_text().strip(), 0)
            
            # Process HTML elements in order
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'pre', 'code', 'blockquote', 'ul', 'ol', 'li']):
                text = element.get_text().strip()
                if not text:
                    continue
                
                try:
                    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        # Headers with proper levels
                        level = min(int(element.name[1]), 9)
                        doc.add_heading(text, level=level)
                    
                    elif element.name in ['pre', 'code']:
                        # Code blocks with monospace font
                        para = doc.add_paragraph()
                        run = para.add_run(text)
                        run.font.name = 'Consolas'
                        run.font.size = 10
                    
                    elif element.name == 'blockquote':
                        # Styled quotes
                        para = doc.add_paragraph(f'"{text}"')
                        para.style = 'Quote'
                    
                    elif element.name == 'li':
                        # List items
                        parent = element.find_parent(['ul', 'ol'])
                        if parent and parent.name == 'ol':
                            doc.add_paragraph(text, style='List Number')
                        else:
                            doc.add_paragraph(text, style='List Bullet')
                    
                    elif element.name in ['p', 'div'] and not element.find_parent(['li', 'blockquote']):
                        # Regular paragraphs (avoid duplicates from nested elements)
                        doc.add_paragraph(text)
                
                except Exception:
                    # Skip problematic elements but continue
                    continue
            
            # Save to BytesIO
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            return {
                "content": output.getvalue(),
                "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "filename": "conversation.docx"
            }
            
        except Exception as e:
            print(f"Enhanced DOCX failed: {e}")
            
            # Fallback to MarkItDown + basic DOCX
            try:
                from docx import Document
                
                # Use MarkItDown for text conversion
                md = MarkItDown(enable_plugins=False)
                input_stream = io.BytesIO(html_bytes)
                result = md.convert_stream(input_stream, source_format="html")
                
                doc = Document()
                doc.add_heading('Converted Document', 0)
                
                # Add content as paragraphs
                for line in result.text_content.splitlines():
                    line = line.strip()
                    if line:
                        if line.startswith('#'):
                            # Simple header detection
                            level = min(len(line) - len(line.lstrip('#')), 9)
                            header_text = line.lstrip('# ')
                            if header_text:
                                doc.add_heading(header_text, level=level)
                        else:
                            doc.add_paragraph(line)
                
                output = io.BytesIO()
                doc.save(output)
                output.seek(0)
                
                return {
                    "content": output.getvalue(),
                    "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "filename": "conversation.docx"
                }
                
            except Exception as final_error:
                print(f"All DOCX methods failed: {final_error}")
                # Return minimal DOCX
                from docx import Document
                doc = Document()
                doc.add_paragraph("Error: Could not convert content to DOCX")
                output = io.BytesIO()
                doc.save(output)
                return {
                    "content": output.getvalue(),
                    "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "filename": "error.docx"
                }
